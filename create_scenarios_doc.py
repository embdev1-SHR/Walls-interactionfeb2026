from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level=1, color=(79, 70, 229)):
    """Add a heading with custom color"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading

def shade_cell(cell, color):
    """Add background color to a cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_heading('Interactive Wall Games & Learning Scenarios', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(79, 70, 229)
title_run.font.size = Pt(24)

# Subtitle
subtitle = doc.add_paragraph('A Comprehensive Guide to Educational Gaming Experiences')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.color.rgb = RGBColor(102, 126, 234)
subtitle_run.font.size = Pt(12)
subtitle_run.italic = True

doc.add_paragraph()

# Table of Contents
doc.add_heading('Contents', level=2)
doc.add_paragraph('1. Creative Games', style='List Number')
doc.add_paragraph('2. Learning Games', style='List Number')
doc.add_paragraph('3. Skill-Building Games', style='List Number')
doc.add_paragraph('4. Interactive Experiences', style='List Number')

doc.add_page_break()

# ===== CREATIVE GAMES SECTION =====
doc.add_heading('1. Creative & Drawing Games', 1)

games_creative = [
    {
        'name': 'Free Canvas',
        'icon': '🎨',
        'description': 'A digital drawing and painting application with multiple brush tools.',
        'tools': ['Pen', 'Soft Brush', 'Marker', 'Spray', 'Ink'],
        'benefits': [
            'Encourages artistic expression and creativity',
            'Develops fine motor skills and hand-eye coordination',
            'Provides a digital canvas for unlimited experimentation',
            'No paper wastage - eco-friendly art creation'
        ]
    },
    {
        'name': 'Coloring Book',
        'icon': '🖍️',
        'description': 'Interactive digital coloring with pre-drawn templates and color palette.',
        'features': ['Multiple coloring templates', 'Color palette selection', 'Undo/Redo functionality'],
        'benefits': [
            'Develops color recognition and selection skills',
            'Improves focus and patience',
            'Creative outlet within structured guidelines',
            'Calming and therapeutic activity'
        ]
    },
    {
        'name': 'Creature Creator',
        'icon': '👾',
        'description': 'Design and customize fantastical creatures with various body parts and features.',
        'features': ['Multiple creature templates', 'Customizable features', 'Save creations'],
        'benefits': [
            'Sparks imagination and creative thinking',
            'Understanding of anatomy and design principles',
            'Storytelling through character creation',
            'Self-expression through visual design'
        ]
    },
    {
        'name': 'Particle Playground',
        'icon': '✨',
        'description': 'Interactive particle effects creation with physics-based animation.',
        'features': ['Touch interactions', 'Real-time particle effects', 'Color customization'],
        'benefits': [
            'Understanding of physics and motion',
            'Visual feedback encourages exploration',
            'Introduces basic animation concepts',
            'Interactive sensory experience'
        ]
    }
]

for game in games_creative:
    add_heading_with_color(doc, f"{game['icon']} {game['name']}", level=2)
    doc.add_paragraph(game['description'])
    
    if 'tools' in game:
        doc.add_paragraph('Available Tools:', style='Heading 4')
        for tool in game['tools']:
            doc.add_paragraph(tool, style='List Bullet')
    
    if 'features' in game:
        doc.add_paragraph('Features:', style='Heading 4')
        for feature in game['features']:
            doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph('How It Helps Kids:', style='Heading 4')
    for benefit in game['benefits']:
        doc.add_paragraph(benefit, style='List Bullet')
    
    # Picture placeholder
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.text = '[Screenshot/Gameplay Video Here]'
    cell_run = cell.paragraphs[0].runs[0]
    cell_run.font.italic = True
    cell_run.font.color.rgb = RGBColor(170, 170, 170)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell, 'E8E8E8')
    
    cell2 = table.rows[1].cells[0]
    cell2.text = 'Picture/Gameplay Screenshot'
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell2.paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()

doc.add_page_break()

# ===== LEARNING GAMES SECTION =====
doc.add_heading('2. Learning & Educational Games', 1)

games_learning = [
    {
        'name': 'Money Shop (Money Counter)',
        'icon': '🛒',
        'category': 'Financial Literacy',
        'description': 'Three-player shopping game where kids learn Indian currency handling and change calculation.',
        'mechanics': ['Tap coins/notes to pay', 'Calculate exact change', 'Multiple difficulty levels'],
        'benefits': [
            'Understanding Indian currency denominations (₹1, ₹2, ₹5, ₹10, ₹20, ₹50, ₹100, ₹200, ₹500)',
            'Mathematical skills: addition and subtraction',
            'Money management and calculation practice',
            'Real-world transaction simulation',
            'Multi-player competitive learning'
        ],
        'curriculum': 'Mathematics - Basic Arithmetic & Money Concepts'
    },
    {
        'name': 'Fruit Counter',
        'icon': '🍎',
        'category': 'Mathematics',
        'description': 'Interactive game for learning numbers through fruit counting mechanics.',
        'mechanics': ['Add/Remove fruits', 'Count and verify', 'Score tracking'],
        'benefits': [
            'Number recognition and sequencing',
            'Basic counting skills development',
            'One-to-one correspondence understanding',
            'Visual representation of quantity',
            'Immediate feedback on accuracy'
        ],
        'curriculum': 'Mathematics - Number Concepts (K-1)'
    },
    {
        'name': 'Number Explorer',
        'icon': '🔢',
        'category': 'Mathematics',
        'description': 'Explore numbers through interactive visual representations and number sequences.',
        'benefits': [
            'Number pattern recognition',
            'Understand skip counting',
            'Prime and composite number concepts',
            'Number relationships and properties',
            'Mathematical thinking development'
        ],
        'curriculum': 'Mathematics - Number Theory'
    },
    {
        'name': 'Word Matcher',
        'icon': '🔤',
        'category': 'Language & Alphabets',
        'description': 'Match words with their meanings through multiple-choice quiz format.',
        'mechanics': ['Multiple choice selection', 'Score tracking', 'Instant feedback'],
        'benefits': [
            'Vocabulary building and expansion',
            'Reading comprehension skills',
            'Word association and memory',
            'Spelling awareness through recognition',
            'Language fluency development'
        ],
        'curriculum': 'English - Vocabulary & Word Recognition'
    },
    {
        'name': 'Alphabet Explorer',
        'icon': '🅰️',
        'category': 'Language & Alphabets',
        'description': 'Interactive alphabet learning with visual representations and sound associations.',
        'benefits': [
            'Alphabet sequence learning',
            'Letter-sound correspondence',
            'Motor skill development through tracing',
            'Phonemic awareness building',
            'Foundation for reading and writing'
        ],
        'curriculum': 'English - Phonics & Alphabet'
    },
    {
        'name': 'Word Creator',
        'icon': '✍️',
        'category': 'Language & Alphabets',
        'description': 'Build words from individual letters with validation and feedback.',
        'benefits': [
            'Spelling practice and improvement',
            'Letter combination understanding',
            'Creative word building',
            'Vocabulary reinforcement',
            'Confidence in word formation'
        ],
        'curriculum': 'English - Spelling & Word Formation'
    },
    {
        'name': 'Arithmetica',
        'icon': '➕',
        'category': 'Mathematics',
        'description': 'Solve arithmetic problems with visual aids and immediate feedback.',
        'benefits': [
            'Addition and subtraction fluency',
            'Mental math development',
            'Problem-solving strategies',
            'Numerical reasoning',
            'Speed and accuracy improvement'
        ],
        'curriculum': 'Mathematics - Basic Arithmetic'
    },
    {
        'name': 'Number Crunch',
        'icon': '🧮',
        'category': 'Mathematics',
        'description': 'Fast-paced number challenges combining multiple mathematical operations.',
        'benefits': [
            'Rapid mental calculation',
            'Multi-operation fluency',
            'Quick decision making',
            'Numeric pattern recognition',
            'Confidence with numbers'
        ],
        'curriculum': 'Mathematics - Multi-operation Arithmetic'
    },
    {
        'name': 'Fruit Math',
        'icon': '🍊',
        'category': 'Mathematics',
        'description': 'Mathematical problems presented through fruit-based scenarios.',
        'benefits': [
            'Real-world math application',
            'Visual learning through graphics',
            'Story problem comprehension',
            'Logical reasoning development',
            'Practical mathematics understanding'
        ],
        'curriculum': 'Mathematics - Applied Arithmetic & Problem Solving'
    },
    {
        'name': 'Word Explorer',
        'icon': '📚',
        'category': 'Language & Alphabets',
        'description': 'Explore words through context, synonyms, antonyms, and definitions.',
        'benefits': [
            'Comprehensive vocabulary expansion',
            'Contextual understanding of words',
            'Synonym and antonym learning',
            'Definition comprehension',
            'Advanced language skills development'
        ],
        'curriculum': 'English - Advanced Vocabulary'
    }
]

for game in games_learning:
    add_heading_with_color(doc, f"{game['icon']} {game['name']}", level=2)
    doc.add_paragraph(f"Category: {game['category']}", style='Heading 4')
    doc.add_paragraph(game['description'])
    
    if 'mechanics' in game:
        doc.add_paragraph('Game Mechanics:', style='Heading 4')
        for mechanic in game['mechanics']:
            doc.add_paragraph(mechanic, style='List Bullet')
    
    doc.add_paragraph('How It Helps Kids:', style='Heading 4')
    for benefit in game['benefits']:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_paragraph('Curriculum Area:', style='Heading 4')
    doc.add_paragraph(game['curriculum'], style='List Bullet')
    
    # Picture placeholder
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.text = '[Screenshot/Gameplay Video Here]'
    cell_run = cell.paragraphs[0].runs[0]
    cell_run.font.italic = True
    cell_run.font.color.rgb = RGBColor(170, 170, 170)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell, 'E8E8E8')
    
    cell2 = table.rows[1].cells[0]
    cell2.text = 'Picture/Gameplay Screenshot'
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell2.paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()

doc.add_page_break()

# ===== SKILL-BUILDING GAMES SECTION =====
doc.add_heading('3. Skill-Building & Life Skills Games', 1)

games_skills = [
    {
        'name': 'Grocery Sorting',
        'icon': '🛒',
        'skill': 'Organization & Categorization',
        'description': 'Sort grocery items into correct categories for organizing shopping.',
        'benefits': [
            'Organization and categorization skills',
            'Understanding food groups and nutrition',
            'Real-world life skill application',
            'Decision-making practice',
            'Cognitive sorting abilities'
        ]
    },
    {
        'name': 'Laundry Sorting',
        'icon': '👕',
        'skill': 'Life Skills & Responsibility',
        'description': 'Sort clothes by color, fabric type, and washing requirements.',
        'benefits': [
            'Household responsibility understanding',
            'Color and pattern recognition',
            'Care and maintenance awareness',
            'Independence and self-care',
            'Practical life skill development'
        ]
    },
    {
        'name': 'Room Cleanup',
        'icon': '🧹',
        'skill': 'Organization & Responsibility',
        'description': 'Organize room items in proper places with visual feedback.',
        'mechanics': ['Bedroom', 'Kitchen', 'Living Room variants'],
        'benefits': [
            'Organization and spatial awareness',
            'Responsibility and ownership',
            'Object categorization',
            'Following instructions',
            'Achievement and task completion'
        ]
    },
    {
        'name': 'Recycling Sort',
        'icon': '♻️',
        'skill': 'Environmental Awareness',
        'description': 'Separate recyclable and non-recyclable waste into proper bins.',
        'benefits': [
            'Environmental consciousness development',
            'Waste management understanding',
            'Sustainability awareness',
            'Responsible citizenship',
            'Category recognition and sorting'
        ]
    },
    {
        'name': 'School Bag Packing',
        'icon': '🎒',
        'skill': 'Organization & Planning',
        'description': 'Pack school bag efficiently with essential items for the day.',
        'benefits': [
            'Planning and time management',
            'Responsibility and independence',
            'Item prioritization',
            'Daily routine understanding',
            'Executive function development'
        ]
    },
    {
        'name': 'Table Setting',
        'icon': '🍽️',
        'skill': 'Etiquette & Social Skills',
        'description': 'Learn proper table setting for meals with correct placement of items.',
        'benefits': [
            'Table etiquette and manners',
            'Social grace and respectful behavior',
            'Spatial arrangement skills',
            'Cultural and family traditions',
            'Host/guest awareness'
        ]
    }
]

for game in games_skills:
    add_heading_with_color(doc, f"{game['icon']} {game['name']}", level=2)
    doc.add_paragraph(f"Skill Focus: {game['skill']}", style='Heading 4')
    doc.add_paragraph(game['description'])
    
    if 'mechanics' in game:
        doc.add_paragraph('Variants:', style='Heading 4')
        for mechanic in game['mechanics']:
            doc.add_paragraph(mechanic, style='List Bullet')
    
    doc.add_paragraph('How It Helps Kids:', style='Heading 4')
    for benefit in game['benefits']:
        doc.add_paragraph(benefit, style='List Bullet')
    
    # Picture placeholder
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.text = '[Screenshot/Gameplay Video Here]'
    cell_run = cell.paragraphs[0].runs[0]
    cell_run.font.italic = True
    cell_run.font.color.rgb = RGBColor(170, 170, 170)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell, 'E8E8E8')
    
    cell2 = table.rows[1].cells[0]
    cell2.text = 'Picture/Gameplay Screenshot'
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell2.paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()

doc.add_page_break()

# ===== INTERACTIVE EXPERIENCES SECTION =====
doc.add_heading('4. Interactive & Exploratory Experiences', 1)

games_interactive = [
    {
        'name': 'Solar System Explorer',
        'icon': '🌍',
        'description': '3D interactive exploration of planets, moons, and celestial objects.',
        'benefits': [
            'Astronomy and space science learning',
            ' 3D spatial understanding',
            'Planetary knowledge acquisition',
            'Science curiosity development',
            'Interactive learning engagement'
        ]
    },
    {
        'name': 'Underwater Explorer',
        'icon': '🐠',
        'description': 'Discover marine life and underwater ecosystems interactively.',
        'benefits': [
            'Marine biology awareness',
            'Ocean ecosystem understanding',
            'Species recognition and knowledge',
            'Environmental conservation awareness',
            'Interactive nature exploration'
        ]
    },
    {
        'name': 'Jungle Games',
        'icon': '🦁',
        'description': 'Explore jungle animals with interactive learning and games.',
        'subgames': ['Forest Animals', 'Hidden Animals'],
        'benefits': [
            'Wildlife knowledge and appreciation',
            'Animal behavior understanding',
            'Biodiversity awareness',
            'Habitat and ecosystem learning',
            'Observation and discovery skills'
        ]
    },
    {
        'name': 'Music Maker',
        'icon': '🎵',
        'description': 'Create music and learn rhythm through interactive sound composition.',
        'benefits': [
            'Musical rhythm and beat understanding',
            'Auditory learning and recognition',
            'Creative sound composition',
            'Music theory basics',
            'Multi-sensory creative expression'
        ]
    },
    {
        'name': 'Rhythm Dance',
        'icon': '💃',
        'description': 'Follow rhythm patterns and dance movements in sync with music.',
        'benefits': [
            'Rhythm and beat coordination',
            'Physical movement and dance',
            'Music-body synchronization',
            'Coordination and motor skills',
            'Fun and energetic activity'
        ]
    },
    {
        'name': '360° Video Player',
        'icon': '🎬',
        'description': 'Immersive 360-degree video experiences for exploratory learning.',
        'benefits': [
            'Immersive virtual reality experience',
            'Exploration of distant places',
            'Cultural and geographical learning',
            'Perspective awareness',
            'Innovative interactive media exposure'
        ]
    },
    {
        'name': 'Butterfly Tracker',
        'icon': '🦋',
        'description': 'Track and identify butterfly species with visual recognition.',
        'benefits': [
            'Entomology and insect knowledge',
            'Pattern and color recognition',
            'Species identification skills',
            'Nature observation',
            'Biodiversity appreciation'
        ]
    }
]

for game in games_interactive:
    add_heading_with_color(doc, f"{game['icon']} {game['name']}", level=2)
    doc.add_paragraph(game['description'])
    
    if 'subgames' in game:
        doc.add_paragraph('Included Games:', style='Heading 4')
        for subgame in game['subgames']:
            doc.add_paragraph(subgame, style='List Bullet')
    
    doc.add_paragraph('How It Helps Kids:', style='Heading 4')
    for benefit in game['benefits']:
        doc.add_paragraph(benefit, style='List Bullet')
    
    # Picture placeholder
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.text = '[Screenshot/Gameplay Video Here]'
    cell_run = cell.paragraphs[0].runs[0]
    cell_run.font.italic = True
    cell_run.font.color.rgb = RGBColor(170, 170, 170)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell, 'E8E8E8')
    
    cell2 = table.rows[1].cells[0]
    cell2.text = 'Picture/Gameplay Screenshot'
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell2.paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()

doc.add_page_break()

# ===== ACTION & SPORTS GAMES SECTION =====
doc.add_heading('5. Action & Sports Games', 1)

games_action = [
    {
        'name': 'Fruit Ninja',
        'icon': '🍎',
        'description': 'Fast-paced slicing game where players cut fruit to earn points.',
        'benefits': [
            'Hand-eye coordination development',
            'Quick reflexes and reaction time',
            'Score and point understanding',
            'Achievement and progression',
            'Engaging action gameplay'
        ]
    },
    {
        'name': 'Goalkeeper Challenge',
        'icon': '⚽',
        'description': 'Sports simulation where players defend goal from incoming shots.',
        'benefits': [
            'Sports skill awareness',
            'Reaction time improvement',
            'Strategic positioning',
            'Achievement motivation',
            'Physical coordination practice'
        ]
    },
    {
        'name': 'Balloon Popper',
        'icon': '🎈',
        'description': 'Pop balloons as they appear on screen with interactive elements.',
        'benefits': [
            'Fine motor skill development',
            'Target accuracy practice',
            'Reflexes and agility',
            'Quick decision making',
            'Simple satisfying gameplay'
        ]
    },
    {
        'name': 'Fish It',
        'icon': '🎣',
        'description': 'Virtual fishing game with catch and collection mechanics.',
        'benefits': [
            'Patience and timing skills',
            'Collection and categorization',
            'Strategy development',
            'Achievement system engagement',
            'Relaxing interactive gameplay'
        ]
    },
    {
        'name': 'Memory Match',
        'icon': '🎮',
        'description': 'Classic memory card matching game for cognitive development.',
        'benefits': [
            'Memory strengthening and recall',
            'Concentration and focus',
            'Cognitive processing speed',
            'Pattern recognition',
            'Brain training and mental exercise'
        ]
    },
    {
        'name': 'Shape Sorter',
        'icon': '🔷',
        'description': 'Sort shapes into matching holes with time and accuracy challenges.',
        'benefits': [
            'Shape recognition and spatial awareness',
            'Problem-solving and logic',
            'Hand-eye coordination',
            'Speed and accuracy balance',
            'Cognitive skill development'
        ]
    },
    {
        'name': 'Puzzle Balance',
        'icon': '⚖️',
        'description': 'Balance weight and elements to solve physics-based puzzles.',
        'benefits': [
            'Physics understanding (weight, balance, gravity)',
            'Problem-solving and strategy',
            'Spatial reasoning',
            'Trial-and-error learning',
            'STEM concept introduction'
        ]
    },
    {
        'name': 'Tug of War',
        'icon': '💪',
        'description': '3D interactive tug of war competition with physics simulation.',
        'benefits': [
            'Competitive gameplay experience',
            'Physics-based interaction understanding',
            'Strategy and timing',
            'Fun multiplayer concept',
            '3D game mechanics exposure'
        ]
    }
]

for game in games_action:
    add_heading_with_color(doc, f"{game['icon']} {game['name']}", level=2)
    doc.add_paragraph(game['description'])
    
    doc.add_paragraph('How It Helps Kids:', style='Heading 4')
    for benefit in game['benefits']:
        doc.add_paragraph(benefit, style='List Bullet')
    
    # Picture placeholder
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.text = '[Screenshot/Gameplay Video Here]'
    cell_run = cell.paragraphs[0].runs[0]
    cell_run.font.italic = True
    cell_run.font.color.rgb = RGBColor(170, 170, 170)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell, 'E8E8E8')
    
    cell2 = table.rows[1].cells[0]
    cell2.text = 'Picture/Gameplay Screenshot'
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell2.paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()

doc.add_page_break()

# ===== SUMMARY SECTION =====
doc.add_heading('Summary & Educational Impact', 1)

summary_text = """This Interactive Wall application provides a comprehensive suite of educational games and experiences designed to support children's learning across multiple domains:

🎨 Creative Expression: Drawing, design, and artistic activities encourage self-expression and creativity.

📚 Academic Learning: Games covering mathematics, language, and literacy help reinforce core academic concepts.

🛠️ Life Skills: Daily task simulations prepare children for practical real-world responsibilities.

🔬 Science & Exploration: Interactive nature and science games foster curiosity and discovery.

🏃 Physical Engagement: Action games develop coordination, reflexes, and physical awareness.

🧠 Cognitive Development: Puzzle games, memory challenges, and brain teasers enhance thinking skills.

Each game is designed with child development principles in mind, providing:
• Immediate feedback and encouragement
• Progressive difficulty and achievement systems
• Multi-sensory learning experiences
• Engaging, gamified learning approaches
• Safe, controlled interactive environments

The application creates a 'learning through play' environment where children can develop academic skills, life skills, and cognitive abilities while enjoying interactive digital experiences."""

doc.add_paragraph(summary_text)

doc.add_paragraph()
doc.add_paragraph('Created for: Interactive Wall Educational Gaming Platform')
doc.add_paragraph(f'Date: February 2026')

# Save document
output_path = r'd:\git\Walls-interactionfeb2026\Interactive_Wall_Games_Scenarios.docx'
doc.save(output_path)
print(f"Document created successfully: {output_path}")
